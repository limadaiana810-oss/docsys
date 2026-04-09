"""
ExportAgent v3 - 按分类独立导出

每种文档类型是独立的导出单元：
  错题集   → 错题.docx（嵌入原图，不加解析）
  经典题   → 经典题.docx（嵌入原图 + LLM 生成解析）
  报销汇总 → 报销.docx（按类型分组汇总）
  通用     → 文档.docx（通用模板）
"""

import json
import re
from pathlib import Path
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
from datetime import datetime

from hub import config, get_writing_llm, extract_json


@dataclass
class ExportResult:
    success: bool
    file_path: Optional[str] = None
    format: str = "docx"
    files_count: int = 0
    error: Optional[str] = None


class ExportAgent:

    def __init__(self):
        self.outbound = Path(config.get("paths.outbound", "/Users/kk/.openclaw/media/outbound/"))
        self.outbound.mkdir(parents=True, exist_ok=True)
        self._llm = None

    def _get_llm(self):
        if self._llm is None:
            self._llm = get_writing_llm()
        return self._llm

    # ── 公开接口 ─────────────────────────────────────────────────────────────

    async def export(
        self,
        results: List[Dict],
        doc_type: str = "general",
        title: str = None,
        format: str = "docx",
    ) -> ExportResult:
        """
        导出主方法。根据 doc_type 选择模板。

        Args:
            results: SearchAgent 返回的结果列表，每个含 storage_path/caption/keywords 等
            doc_type: "wrong_questions" / "classic_questions" / "reimbursement" / "general"
            title: 自定义标题（可选）
            format: "docx" / "pdf"
        """
        if not results:
            return ExportResult(success=False, error="没有要导出的内容")

        try:
            if doc_type == "wrong_questions":
                result = await self._export_wrong_questions(results, title)
            elif doc_type == "classic_questions":
                result = await self._export_classic_questions(results, title)
            elif doc_type == "reimbursement":
                result = await self._export_reimbursement(results, title)
            else:
                result = await self._export_general(results, title)

            if result.success and format == "pdf":
                pdf_result = self._convert_to_pdf(result.file_path)
                if pdf_result:
                    return ExportResult(success=True, file_path=pdf_result, format="pdf",
                                        files_count=result.files_count)

            return result

        except Exception as e:
            return ExportResult(success=False, error=str(e))

    # ── 错题集：严格复刻原题，嵌入原图 ──────────────────────────────────────

    async def _export_wrong_questions(
        self, results: List[Dict], title: str = None
    ) -> ExportResult:
        """
        错题.docx — 只嵌入原图，不加分析不加解法。
        按学科分组，每题标注日期和关键词。
        """
        doc, _set_font, _add_line = self._init_doc()
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        title = title or "错题集"
        self._add_title(doc, title, _set_font)

        # 按学科分组
        groups = self._group_by_subject(results)

        total = 0
        for subject, items in groups.items():
            # 学科标题
            self._add_section_heading(doc, f"{subject}（{len(items)} 题）", _set_font)

            for idx, item in enumerate(items, 1):
                img_path = item.get("storage_path", "")
                caption = item.get("caption", "")
                keywords = item.get("keywords", [])
                if isinstance(keywords, str):
                    keywords = [k.strip() for k in keywords.split(",") if k.strip()]

                # 题号
                p = doc.add_paragraph()
                r = p.add_run(f"第 {idx} 题")
                _set_font(r, size_pt=11, bold=True)
                if keywords:
                    r2 = p.add_run(f"  [{', '.join(keywords[:4])}]")
                    _set_font(r2, size_pt=9, color=RGBColor(120, 120, 120))

                # 嵌入原图
                if img_path and Path(img_path).exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception:
                        doc.add_paragraph(f"[图片加载失败: {Path(img_path).name}]")
                else:
                    doc.add_paragraph(f"[文件不存在: {img_path}]")

                doc.add_paragraph()  # 间距
                total += 1

            _add_line()

        return self._save_doc(doc, title, total)

    # ── 经典题：原图 + 解析 ──────────────────────────────────────────────────

    async def _export_classic_questions(
        self, results: List[Dict], title: str = None
    ) -> ExportResult:
        """
        经典题.docx — 嵌入原图 + LLM 生成每题解析。
        """
        doc, _set_font, _add_line = self._init_doc()
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        title = title or "经典题集"
        self._add_title(doc, title, _set_font)

        # 批量生成解析
        analyses = await self._batch_analyze(results)

        groups = self._group_by_subject(results)

        total = 0
        for subject, items in groups.items():
            self._add_section_heading(doc, f"{subject}（{len(items)} 题）", _set_font)

            for idx, item in enumerate(items, 1):
                img_path = item.get("storage_path", "")
                record_id = item.get("record_id", "")
                keywords = item.get("keywords", [])
                if isinstance(keywords, str):
                    keywords = [k.strip() for k in keywords.split(",") if k.strip()]

                # 题号 + 知识点
                p = doc.add_paragraph()
                r = p.add_run(f"第 {idx} 题")
                _set_font(r, size_pt=11, bold=True)
                if keywords:
                    r2 = p.add_run(f"  [{', '.join(keywords[:4])}]")
                    _set_font(r2, size_pt=9, color=RGBColor(80, 80, 160))

                # 嵌入原图
                if img_path and Path(img_path).exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception:
                        doc.add_paragraph(f"[图片加载失败: {Path(img_path).name}]")

                # 解析
                analysis = analyses.get(record_id, "")
                if analysis:
                    ap = doc.add_paragraph()
                    ar = ap.add_run("📝 解析：")
                    _set_font(ar, size_pt=10, bold=True, color=RGBColor(0, 100, 0))
                    ar2 = ap.add_run(f" {analysis}")
                    _set_font(ar2, size_pt=10)

                doc.add_paragraph()
                total += 1

            _add_line()

        return self._save_doc(doc, title, total)

    async def _batch_analyze(self, results: List[Dict]) -> Dict[str, str]:
        """为经典题批量生成解析"""
        llm = self._get_llm()
        analyses = {}

        # 构建批量 prompt
        items_text = []
        for i, r in enumerate(results[:20]):
            caption = r.get("caption", "")
            extracted = r.get("extracted_text", "")
            text = extracted if extracted else caption
            items_text.append(f"[{r.get('record_id', i)}] {text[:200]}")

        if not items_text:
            return analyses

        system = """你是一个优秀的学科老师。为每道题生成简洁的解析（1-3句话），包含：
1. 解题思路
2. 关键步骤
3. 易错点提醒

输出 JSON 对象，key 是题目 ID，value 是解析文字。
只输出 JSON。"""

        try:
            raw = await llm.generate(
                "\n".join(items_text), system=system, max_tokens=2048
            )
            data = extract_json(raw)
            if isinstance(data, dict):
                analyses = {str(k): str(v) for k, v in data.items()}
        except Exception as e:
            print(f"   ⚠️ 批量解析失败: {e}")

        return analyses

    # ── 报销汇总 ─────────────────────────────────────────────────────────────

    async def _export_reimbursement(
        self, results: List[Dict], title: str = None
    ) -> ExportResult:
        """报销.docx — 按报销类型分组汇总"""
        doc, _set_font, _add_line = self._init_doc()
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        title = title or "报销汇总"
        self._add_title(doc, title, _set_font)

        # 按 keywords 中的类型分组
        groups = {}
        for r in results:
            keywords = r.get("keywords", [])
            if isinstance(keywords, str):
                keywords = [k.strip() for k in keywords.split(",")]
            category = "其他"
            for kw in keywords:
                if kw in ("餐饮", "交通", "差旅", "办公", "住宿"):
                    category = kw
                    break
            groups.setdefault(category, []).append(r)

        total = 0
        for category, items in groups.items():
            self._add_section_heading(doc, f"{category}（{len(items)} 张）", _set_font)

            for idx, item in enumerate(items, 1):
                img_path = item.get("storage_path", "")
                caption = item.get("caption", "")

                p = doc.add_paragraph()
                r = p.add_run(f"{idx}. {caption}")
                _set_font(r, size_pt=10)

                if img_path and Path(img_path).exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(4.5))
                        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception:
                        pass

                doc.add_paragraph()
                total += 1

            _add_line()

        return self._save_doc(doc, title, total)

    # ── 通用导出 ─────────────────────────────────────────────────────────────

    async def _export_general(
        self, results: List[Dict], title: str = None
    ) -> ExportResult:
        """通用文档导出"""
        doc, _set_font, _add_line = self._init_doc()
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        title = title or "文档导出"
        self._add_title(doc, title, _set_font)

        for idx, item in enumerate(results, 1):
            img_path = item.get("storage_path", "")
            caption = item.get("caption", "")

            p = doc.add_paragraph()
            r = p.add_run(f"{idx}. {caption}")
            _set_font(r, size_pt=10)

            if img_path and Path(img_path).exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(5.0))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception:
                    pass

            doc.add_paragraph()

        return self._save_doc(doc, title, len(results))

    # ── 共享工具方法 ─────────────────────────────────────────────────────────

    def _init_doc(self):
        """初始化 Word 文档 + 返回辅助函数"""
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise RuntimeError("python-docx 未安装，请运行: pip install python-docx")

        doc = Document()

        for sec in doc.sections:
            sec.page_width = Cm(21)
            sec.page_height = Cm(29.7)
            sec.left_margin = Cm(2)
            sec.right_margin = Cm(2)
            sec.top_margin = Cm(2)
            sec.bottom_margin = Cm(2)

        def _set_font(run, size_pt=11, bold=False, color=None):
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            run.font.name = "微软雅黑"
            run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if color:
                run.font.color.rgb = color

        def _add_line():
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)

        return doc, _set_font, _add_line

    def _add_title(self, doc, title: str, _set_font):
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import RGBColor

        h = doc.add_heading("", 0)
        r = h.add_run(title)
        _set_font(r, size_pt=18, bold=True)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        h.paragraph_format.space_after = Pt(6)

        date_p = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日')}")
        date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if date_p.runs:
            _set_font(date_p.runs[0], size_pt=9, color=RGBColor(150, 150, 150))

    def _add_section_heading(self, doc, text: str, _set_font):
        from docx.shared import Pt
        h = doc.add_heading("", 1)
        r = h.add_run(text)
        _set_font(r, size_pt=14, bold=True)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

    def _group_by_subject(self, results: List[Dict]) -> Dict[str, List[Dict]]:
        """按学科分组，无学科的归入'综合'"""
        subject_tags = {"语文", "数学", "英语", "物理", "化学", "生物", "历史", "地理", "道德与法治", "音乐", "美术", "体育", "信息技术"}
        groups = {}
        for r in results:
            # 优先从 category 字段获取学科
            category = r.get("category", "")
            if category and category in subject_tags:
                subject = category
            else:
                # fallback: 从 keywords 中查找
                keywords = r.get("keywords", [])
                if isinstance(keywords, str):
                    keywords = [k.strip() for k in keywords.split(",")]
                subject = "综合"
                for kw in keywords:
                    if kw in subject_tags:
                        subject = kw
                        break
            groups.setdefault(subject, []).append(r)
        return groups

    def _save_doc(self, doc, title: str, count: int) -> ExportResult:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)
        filename = f"{safe_title}_{ts}.docx"
        output_path = self.outbound / filename
        doc.save(str(output_path))
        print(f"   ✅ 导出完成: {output_path} ({count} 条)")
        return ExportResult(success=True, file_path=str(output_path), format="docx",
                            files_count=count)

    def _convert_to_pdf(self, docx_path: str) -> Optional[str]:
        """尝试将 docx 转为 pdf"""
        import subprocess
        import shutil

        docx = Path(docx_path)
        pdf = docx.with_suffix(".pdf")

        for cmd in [shutil.which("libreoffice"), shutil.which("soffice")]:
            if cmd:
                try:
                    subprocess.run(
                        [cmd, "--headless", "--convert-to", "pdf",
                         "--outdir", str(docx.parent), str(docx)],
                        check=True, capture_output=True, timeout=60,
                    )
                    docx.unlink(missing_ok=True)
                    return str(pdf)
                except Exception:
                    pass

        if shutil.which("pandoc"):
            try:
                subprocess.run(
                    ["pandoc", str(docx), "-o", str(pdf)],
                    check=True, capture_output=True, timeout=60,
                )
                docx.unlink(missing_ok=True)
                return str(pdf)
            except Exception:
                pass

        return None
